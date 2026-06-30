import base64
import logging
from io import BytesIO

import httpx
import pdfplumber
from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from app.config import settings

logger = logging.getLogger(__name__)


class ParserService:
    """从上传的 PRD 文件中提取文本，支持 PDF、DOCX、MD、TXT 和图片识别。"""

    @staticmethod
    async def parse(filename: str, content: bytes) -> str:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

        if ext == "pdf":
            return ParserService._parse_pdf(content)
        elif ext == "docx":
            return await ParserService._parse_docx(content)
        elif ext in ("md", "txt"):
            return content.decode("utf-8", errors="replace")
        else:
            try:
                return content.decode("utf-8")
            except UnicodeDecodeError:
                return f"[不支持的文件格式: {ext}]"

    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        text_parts = []
        with pdfplumber.open(BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n\n".join(text_parts) if text_parts else "[PDF 提取内容为空]"

    @staticmethod
    async def _parse_docx(content: bytes) -> str:
        doc = DocxDocument(BytesIO(content))
        parts = []

        # 提取段落文本。
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # 提取表格文本。
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        # 通过多模态 LLM 提取并描述图片内容。
        image_descriptions = await ParserService._extract_images(doc)
        if image_descriptions:
            parts.append("\n## 文档中的图片/截图内容\n" + image_descriptions)

        return "\n".join(parts) if parts else "[Word 文档提取内容为空]"

    @staticmethod
    async def _extract_images(doc: DocxDocument) -> str:
        """提取 docx 内图片并用多模态 LLM 描述。"""
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    # 识别图片格式。
                    ext = image_part.partname.split(".")[-1].lower() if "." in image_part.partname else "png"
                    mime_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "gif": "image/gif", "webp": "image/webp", "bmp": "image/bmp"}
                    mime = mime_map.get(ext, "image/png")
                    b64 = base64.b64encode(image_bytes).decode()
                    images.append({"data": b64, "mime": mime})
                except Exception:
                    logger.exception("提取 Word 图片失败，跳过当前图片")
                    continue

        if not images:
            return ""

        # 使用多模态 LLM 描述图片。
        descriptions = await ParserService._describe_images(images)
        return descriptions

    @staticmethod
    async def _describe_images(images: list[dict]) -> str:
        """把图片发送给多模态 LLM 获取描述。"""
        if not settings.LLM_API_KEY:
            return "（未配置API Key，无法识别图片）"

        # 构造包含文本和图片的多模态消息。
        content_parts = [{
            "type": "text",
            "text": "请逐一描述以下图片中的内容，特别是UI界面、表格数据、流程图、架构图等对测试用例设计有帮助的信息。用中文输出，每张图片的描述不超过200字。"
        }]

        for img in images[:5]:  # 最多识别 5 张图片
            content_parts.append({
                "type": "image_url",
                "image_url": {"url": f"data:{img['mime']};base64,{img['data']}"}
            })

        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    f"{settings.LLM_BASE_URL}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {settings.LLM_API_KEY}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model": settings.LLM_MODEL,
                        "messages": [{"role": "user", "content": content_parts}],
                        "max_tokens": 2048,
                    },
                )
                response.raise_for_status()
                data = response.json()
                return data["choices"][0]["message"]["content"]
        except Exception as e:
            logger.exception("图片识别请求失败")
            return f"[图片识别失败: {str(e)}]"
